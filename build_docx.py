#!/usr/bin/env python3
"""Generate a Springer-formatted .docx manuscript from book_chapter_final.md.

- Springer-style typography (Times New Roman, numbered headings, justified body).
- Embeds outputs/figures/*.png with "Fig. N" captions at mapped sections.
- Booktabs-style tables (no vertical rules; rules above/below header and at bottom).
- Numbered [n] citations resolved from outputs/literature_matrix_updated.csv,
  with a generated numbered reference list (cited-only, first-appearance order).
"""
import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

ROOT = Path(__file__).resolve().parent
MD = ROOT / "outputs" / "chapter" / "book_chapter_final.md"
FIGDIR = ROOT / "outputs" / "figures"
LITCSV = ROOT / "outputs" / "literature_matrix_updated.csv"
OUT = ROOT / "Springer_Manuscript.docx"

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"

# Figures inserted at the END of the section whose heading contains the key text.
FIG_MAP = {
    "Conceptual definition": [
        ("framework_diagram.png",
         "Scientific workflow of the Concrete Performance Passport: from raw "
         "concrete data and ML prediction through the uncertainty/risk/evidence "
         "layer to OpenBIM/IDS information requirements and QA/QC and "
         "sustainability decision support."),
    ],
    "6.1 Compressive-strength model performance": [
        ("strength_model_comparison.png",
         "Held-out coefficient of determination (R\u00b2) and RMSE for the compared "
         "compressive-strength models. Gradient-boosted trees lead the linear baseline."),
    ],
    "6.2 Explainability": [
        ("strength_feature_importance.png",
         "Impurity-based feature importance for the best compressive-strength model."),
        ("strength_explainability.png",
         "Permutation feature importance for the best compressive-strength model; "
         "age and cement content dominate."),
    ],
    "6.3 Carbon benchmarking": [
        ("carbon_distribution.png",
         "Distribution of A1\u2013A3 global warming potential (GWP) across the public "
         "EPD dataset (n = 43,053) and by compressive-strength class."),
    ],
    "6.4 Passport demonstration": [
        ("risk_class_distribution.png",
         "Distribution of risk classes and QA/QC decisions across the demonstration "
         "Concrete Performance Passports."),
    ],
}

# ---------------------------------------------------------------------------
# Citations
# ---------------------------------------------------------------------------
def load_refs():
    refs = {}
    with open(LITCSV, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            key = (row.get("citation_key") or "").strip()
            if key:
                refs[key] = row
    return refs


REFS = load_refs()
_cite_order = []  # citation keys in first-appearance order


def _cite_num(key):
    if key not in _cite_order:
        _cite_order.append(key)
    return _cite_order.index(key) + 1


def resolve_citations(text):
    """Replace backtick tokens that are known citation keys with [n]; merge runs."""
    def repl(m):
        tok = m.group(1)
        if tok in REFS:
            return f"[{_cite_num(tok)}]"
        return m.group(0)  # leave non-citation code spans untouched

    text = re.sub(r"`([^`]+)`", repl, text)
    # Merge "[1], [2], [3]" / "[1]; [2]" -> "[1, 2, 3]"
    def merge(m):
        nums = re.findall(r"\d+", m.group(0))
        return "[" + ", ".join(nums) + "]"
    text = re.sub(r"\[\d+\](?:[;,]\s*\[\d+\])+", merge, text)
    # "([1])" -> "[1]"
    text = re.sub(r"\(\[(\d+(?:, \d+)*)\]\)", r"[\1]", text)
    return text


def format_reference(n, row):
    authors = (row.get("authors") or "").strip()
    if authors.startswith("(") or not authors:
        authors = ""
    year = (row.get("year") or "").strip()
    title = (row.get("title") or "").strip()
    venue = (row.get("venue") or "").strip()
    publisher = (row.get("publisher") or "").strip()
    doi = (row.get("doi") or "").strip()
    url = (row.get("url") or "").strip()
    parts = []
    if authors:
        parts.append(f"{authors}")
    if title:
        parts.append(title)
    tail = venue
    if publisher and publisher.lower() not in venue.lower():
        tail = f"{tail}. {publisher}" if tail else publisher
    if tail:
        parts.append(tail)
    if year:
        parts.append(f"({year})")
    parts = [p.rstrip(". ").strip() for p in parts if p and p.strip()]
    s = ". ".join(parts)
    if doi:
        s += f". https://doi.org/{doi}" if not doi.startswith("http") else f". {doi}"
    elif url:
        s += f". {url}"
    return s


# ---------------------------------------------------------------------------
# DOCX low-level helpers
# ---------------------------------------------------------------------------
def set_cell_border(cell, edges):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, sz in edges.items():
        el = tcBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tcBorders.append(el)
        if sz == 0:
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")


def add_runs_with_markup(paragraph, text, base_size=10):
    """Render **bold**, *italic*, and `code` inline markup into runs."""
    # tokenizer for **bold**, *italic*, `code`
    pattern = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.size = Pt(base_size)
        tok = m.group(0)
        if tok.startswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*"):
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        else:  # code
            r = paragraph.add_run(tok[1:-1]); r.font.name = MONO_FONT
            r.font.size = Pt(base_size - 1)
        r.font.size = r.font.size or Pt(base_size)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:]); r.font.size = Pt(base_size)


def strip_artifact_parens(text):
    # Remove parenthetical cross-references to repo artefacts, e.g. (`outputs/...`)
    text = re.sub(r"\s*\((?:see\s+)?`?outputs/[^)]*\)", "", text)
    text = re.sub(r"\s*\(`?[^)]*\.png`?\)", "", text)
    return text


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------
def parse_blocks(md_text):
    """Yield structural blocks: ('h0'|'h1'|'h2', text), ('p', text),
    ('ul'|'ol', [items]), ('table', (caption, header, rows)), ('quote', text)."""
    lines = md_text.split("\n")
    i = 0
    pending_table_caption = None
    while i < len(lines):
        line = lines[i]
        s = line.rstrip("\n")
        if not s.strip():
            i += 1
            continue
        if s.startswith("# "):
            yield ("h0", s[2:].strip()); i += 1; continue
        if s.startswith("## "):
            yield ("h1", s[3:].strip()); i += 1; continue
        if s.startswith("### "):
            yield ("h2", s[4:].strip()); i += 1; continue
        if s.strip() == "---":
            i += 1; continue
        if s.startswith(">"):
            # gather blockquote
            buf = []
            while i < len(lines) and lines[i].startswith(">"):
                buf.append(lines[i].lstrip(">").strip()); i += 1
            yield ("quote", " ".join(x for x in buf if x)); continue
        # table caption like **Table 1.** ...
        mcap = re.match(r"^\*\*Table\s+\d+\.\*\*", s)
        if mcap:
            pending_table_caption = s
            i += 1
            # skip optional blank then expect table
            while i < len(lines) and not lines[i].strip():
                i += 1
            continue
        if s.lstrip().startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                tbl.append(lines[i].strip()); i += 1
            header = [c.strip() for c in tbl[0].strip("|").split("|")]
            rows = []
            for r in tbl[2:]:
                rows.append([c.strip() for c in r.strip("|").split("|")])
            yield ("table", (pending_table_caption, header, rows))
            pending_table_caption = None
            continue
        if re.match(r"^\d+\.\s+", s):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip())); i += 1
            yield ("ol", items); continue
        if s.lstrip().startswith("- "):
            items = []
            while i < len(lines) and lines[i].lstrip().startswith("- "):
                items.append(lines[i].lstrip()[2:].strip()); i += 1
            yield ("ul", items); continue
        # default paragraph (single line; md uses blank-line separation)
        yield ("p", s.strip()); i += 1


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
def build():
    md_text = MD.read_text(encoding="utf-8")
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.0
    # East-asian font binding
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    # Page setup (Springer-ish)
    sec = doc.sections[0]
    sec.page_height = Inches(9.61); sec.page_width = Inches(6.69)  # ~ Springer 17x24.4cm
    sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)

    fig_counter = {"n": 0}
    table_counter = {"n": 0}
    current_section_title = ""
    pending_section_figs = []

    def flush_section_figures():
        # insert any figures mapped to the section we are leaving
        for fname, caption in pending_section_figs:
            fpath = FIGDIR / fname
            if not fpath.exists():
                continue
            fig_counter["n"] += 1
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(fpath), width=Inches(4.9))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(2)
            r = cap.add_run(f"Fig. {fig_counter['n']} ")
            r.bold = True; r.font.size = Pt(9)
            add_runs_with_markup(cap, caption, base_size=9)
        pending_section_figs.clear()

    def figs_for(title):
        for key, figs in FIG_MAP.items():
            if key.lower() in title.lower():
                return list(figs)
        return []

    blocks = list(parse_blocks(md_text))
    in_references = False

    for kind, payload in blocks:
        # When we hit a new heading, flush figures for the previous section.
        if kind in ("h1", "h2", "h0"):
            flush_section_figures()

        if kind == "h0":
            h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = h.add_run(payload); r.bold = True; r.font.size = Pt(15)
            h.paragraph_format.space_after = Pt(10)
            # Author/affiliation placeholder block (fill before submission)
            a = doc.add_paragraph()
            ar = a.add_run("\u27e8Author name(s)\u27e9")
            ar.italic = True; ar.font.size = Pt(11)
            aff = doc.add_paragraph()
            afr = aff.add_run("\u27e8Affiliation, city, country\u27e9  \u00b7  \u27e8email@domain\u27e9")
            afr.italic = True; afr.font.size = Pt(9)
            aff.paragraph_format.space_after = Pt(10)
            continue

        if kind == "h1":
            in_references = payload.lower().startswith("references")
            current_section_title = payload
            pending_section_figs = figs_for(payload)
            h = doc.add_heading(level=1)
            r = h.add_run(payload); r.bold = True
            r.font.name = BODY_FONT; r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
            if in_references:
                # Render generated numbered references then stop normal flow
                _render_references(doc)
            continue

        if kind == "h2":
            current_section_title = payload
            pending_section_figs = figs_for(payload)
            h = doc.add_heading(level=2)
            r = h.add_run(payload); r.bold = True; r.italic = True
            r.font.name = BODY_FONT; r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0, 0, 0)
            continue

        if in_references:
            # skip the original placeholder paragraph under References
            continue

        if kind == "p":
            text = strip_artifact_parens(payload)
            text = resolve_citations(text)
            # Abstract/Keywords lead-in styling
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs_with_markup(p, text, base_size=10)
            continue

        if kind == "quote":
            text = strip_artifact_parens(payload)
            if "figure" in text.lower():
                continue  # figures handled via mapping
            # Render meta-notes verbatim (do NOT resolve citations here, so an
            # example key like `broyles2024epd` stays literal and numbering
            # starts in the body at [1]).
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs_with_markup(p, text, base_size=9)
            for run in p.runs:
                run.italic = True
            continue

        if kind == "ul":
            for it in payload:
                it = resolve_citations(strip_artifact_parens(it))
                p = doc.add_paragraph(style="List Bullet")
                add_runs_with_markup(p, it, base_size=10)
            continue

        if kind == "ol":
            for it in payload:
                it = resolve_citations(strip_artifact_parens(it))
                p = doc.add_paragraph(style="List Number")
                add_runs_with_markup(p, it, base_size=10)
            continue

        if kind == "table":
            caption, header, rows = payload
            table_counter["n"] += 1
            cap = doc.add_paragraph()
            cap.paragraph_format.space_before = Pt(6)
            r = cap.add_run(f"Table {table_counter['n']} ")
            r.bold = True; r.font.size = Pt(9)
            if caption:
                ctext = re.sub(r"^\*\*Table\s+\d+\.\*\*\s*", "", caption)
                ctext = strip_artifact_parens(resolve_citations(ctext))
                add_runs_with_markup(cap, ctext, base_size=9)
            _render_table(doc, header, rows)
            continue

    flush_section_figures()
    doc.save(str(OUT))
    print("Saved", OUT)
    print("Figures embedded:", fig_counter["n"], "| Tables:", table_counter["n"],
          "| Citations:", len(_cite_order))


def _render_table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    # header
    hdr = t.rows[0].cells
    for j, htext in enumerate(header):
        p = hdr[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
        add_runs_with_markup(p, htext, base_size=9)
        for run in p.runs:
            run.bold = True
    # data rows
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_runs_with_markup(p, val, base_size=9)
    # booktabs borders: top of header, bottom of header, bottom of last row; no verticals
    ncols = len(header)
    for j in range(ncols):
        set_cell_border(t.rows[0].cells[j], {"top": 8, "bottom": 6,
                                             "left": 0, "right": 0, "insideH": 0, "insideV": 0})
    last = t.rows[-1].cells
    for j in range(ncols):
        set_cell_border(last[j], {"bottom": 8, "left": 0, "right": 0, "insideV": 0})
    # ensure no vertical/internal lines on middle rows
    for ri in range(1, len(t.rows) - 1):
        for j in range(ncols):
            set_cell_border(t.rows[ri].cells[j], {"left": 0, "right": 0,
                                                  "top": 0, "bottom": 0, "insideV": 0})


def _render_references(doc):
    """Render the numbered reference list for all cited keys (first-appearance order).

    Called when the 'References' heading is reached; by then every in-text
    citation has already been resolved, so _cite_order is complete.
    """
    if not _cite_order:
        return
    for i, key in enumerate(_cite_order, start=1):
        row = REFS.get(key, {})
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(3)
        rn = p.add_run(f"{i}. ")
        rn.bold = True; rn.font.size = Pt(9)
        rr = p.add_run(format_reference(i, row))
        rr.font.size = Pt(9)


if __name__ == "__main__":
    build()
